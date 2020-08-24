import docx

name_list = open('guests.txt').readlines()

doc = docx.Document()
for name in name_list:
    name=name.strip('\n')
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.add_paragraph(name)
    doc.add_paragraph('At 11010 memory at the evening of')
    doc.add_paragraph('April 1st')
    doc.add_paragraph("At 7 o'clock")
    doc.add_page_break()
doc.save('邀请函.docx')
print('Done')
