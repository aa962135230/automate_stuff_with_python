import docx

doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('zophie.docx')